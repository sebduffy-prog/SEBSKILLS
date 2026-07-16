#!/usr/bin/env python3
"""Produce a Word tracked-changes (redline) document from two .docx versions.

Aligns paragraphs between OLD and NEW, then word-diffs each aligned pair,
emitting native <w:ins>/<w:del> runs that Word/Google Docs render as tracked
changes (Review > All Markup). Deleted text uses <w:delText>; inserted text
uses <w:t>. No native python-docx API exists for this — we build the OOXML.

Usage:
    python3 redline_compare.py OLD.docx NEW.docx OUT.docx [--author "Name"] [--char]

--char : diff at character granularity inside a paragraph (default: word).
"""
import argparse
import datetime as dt
import difflib
import sys

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _run(text, deleted=False):
    r = OxmlElement("w:r")
    t = OxmlElement("w:delText" if deleted else "w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _wrap(tag, author, date, wid, child):
    el = OxmlElement(tag)  # "w:ins" or "w:del"
    el.set(qn("w:id"), str(wid))
    el.set(qn("w:author"), author)
    el.set(qn("w:date"), date)
    el.append(child)
    return el


def _para_text(p):
    return p.text


def _tokens(s, char_mode):
    if char_mode:
        return list(s)
    # keep trailing spaces attached so re-joins read naturally
    return [w + " " for w in s.split(" ")] if s else []


def _join(tokens, char_mode):
    return "".join(tokens) if char_mode else "".join(tokens)


class _Counter:
    def __init__(self):
        self.n = 0

    def next(self):
        self.n += 1
        return self.n


def _diff_into_paragraph(p, old_s, new_s, author, date, ids, char_mode):
    a, b = _tokens(old_s, char_mode), _tokens(new_s, char_mode)
    sm = difflib.SequenceMatcher(None, a, b, autojunk=False)
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            p._p.append(_run(_join(a[i1:i2], char_mode)))
        if tag in ("replace", "delete"):
            txt = _join(a[i1:i2], char_mode)
            if txt:
                p._p.append(_wrap("w:del", author, date, ids.next(),
                                  _run(txt, deleted=True)))
        if tag in ("replace", "insert"):
            txt = _join(b[j1:j2], char_mode)
            if txt:
                p._p.append(_wrap("w:ins", author, date, ids.next(), _run(txt)))


def build_redline(old_path, new_path, out_path, author, char_mode):
    old_paras = [_para_text(p) for p in Document(old_path).paragraphs]
    new_paras = [_para_text(p) for p in Document(new_path).paragraphs]
    date = dt.datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")
    ids = _Counter()

    out = Document()
    # drop any default empty paragraph so output starts clean
    for p in list(out.paragraphs):
        p._p.getparent().remove(p._p)

    sm = difflib.SequenceMatcher(None, old_paras, new_paras, autojunk=False)
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            for k in range(i1, i2):
                p = out.add_paragraph()
                p._p.append(_run(old_paras[k]))
        elif tag == "replace":
            # align replaced blocks pairwise; word-diff each pair
            span = max(i2 - i1, j2 - j1)
            for off in range(span):
                oi, nj = i1 + off, j1 + off
                p = out.add_paragraph()
                o = old_paras[oi] if oi < i2 else ""
                n = new_paras[nj] if nj < j2 else ""
                _diff_into_paragraph(p, o, n, author, date, ids, char_mode)
        elif tag == "delete":
            for k in range(i1, i2):
                p = out.add_paragraph()
                _diff_into_paragraph(p, old_paras[k], "", author, date, ids, char_mode)
        elif tag == "insert":
            for k in range(j1, j2):
                p = out.add_paragraph()
                _diff_into_paragraph(p, "", new_paras[k], author, date, ids, char_mode)

    out.save(out_path)
    return ids.n


def main():
    ap = argparse.ArgumentParser(description="Build a tracked-changes redline from two .docx files.")
    ap.add_argument("old"); ap.add_argument("new"); ap.add_argument("out")
    ap.add_argument("--author", default="Redline Compare")
    ap.add_argument("--char", action="store_true", help="character-level diff")
    args = ap.parse_args()
    n = build_redline(args.old, args.new, args.out, args.author, args.char)
    print(f"Wrote {args.out} with {n} tracked change(s), author={args.author!r}")


if __name__ == "__main__":
    sys.exit(main())
